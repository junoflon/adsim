"""
다양한 파일 포맷에서 텍스트 추출
- PDF (.pdf) — PyMuPDF
- DOCX (.docx) — python-docx
- HWP (.hwp) — olefile (한글 5.x 바이너리)
- HWPX (.hwpx) — zip 내부 content.xml 파싱
- TXT / MD — 원문 인코딩 감지
"""

import os
import re
import zipfile
from typing import Optional


def extract_text(file_path: str, max_chars: int = 20000) -> str:
    """파일 확장자에 따라 텍스트 추출. 실패 시 빈 문자열 반환."""
    if not file_path or not os.path.exists(file_path):
        return ""
    ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""
    try:
        if ext == "pdf":
            text = _from_pdf(file_path)
        elif ext == "docx":
            text = _from_docx(file_path)
        elif ext == "hwp":
            text = _from_hwp(file_path)
        elif ext == "hwpx":
            text = _from_hwpx(file_path)
        elif ext == "xlsx":
            text = _from_xlsx(file_path)
        elif ext == "csv":
            text = _from_csv(file_path)
        elif ext in ("txt", "md"):
            text = _from_text(file_path)
        else:
            text = ""
    except Exception as e:
        text = f"(파일 읽기 실패: {e})"

    text = _clean_whitespace(text)
    return text[:max_chars]


def _from_pdf(path: str) -> str:
    import fitz  # type: ignore
    doc = fitz.open(path)
    try:
        parts = []
        for page in doc:
            parts.append(page.get_text("text"))
        return "\n".join(parts)
    finally:
        doc.close()


def _from_docx(path: str) -> str:
    import docx  # type: ignore
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 표 내용도 읽기
    for table in d.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _from_hwp(path: str) -> str:
    """HWP 5.x 이하 바이너리. olefile로 BodyText/Section* 스트림을 읽어 텍스트만 추출."""
    import olefile  # type: ignore
    ole = olefile.OleFileIO(path)
    try:
        streams = [entry for entry in ole.listdir() if entry[0] == "BodyText"]
        parts = []
        for path_parts in streams:
            try:
                data = ole.openstream("/".join(path_parts)).read()
                # HWP 문자열은 UTF-16 LE. 인쇄 가능한 한글/영문/숫자만 필터.
                decoded = data.decode("utf-16-le", errors="ignore")
                # 제어문자/깨진 문자 제거
                cleaned = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]", " ", decoded)
                # 한글 2글자 이상 연속되거나 영문 단어 단위만 유지
                tokens = re.findall(r"[가-힣]{2,}[가-힣\s,.!?~%()·\-]*|[A-Za-z][A-Za-z0-9\s,.!?\-]*|\d+[%가-힣]*", cleaned)
                parts.extend(t.strip() for t in tokens if len(t.strip()) > 1)
            except Exception:
                continue
        return "\n".join(parts)
    finally:
        ole.close()


def _from_hwpx(path: str) -> str:
    """HWPX는 zip 포맷. Contents/section0.xml 등에서 <hp:t> 태그의 텍스트 추출."""
    parts = []
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if name.startswith("Contents/section") and name.endswith(".xml"):
                try:
                    xml = z.read(name).decode("utf-8", errors="ignore")
                    # <hp:t>텍스트</hp:t> 추출 — XML 파서 대신 regex (의존성 최소화)
                    matches = re.findall(r"<hp:t[^>]*>([^<]*)</hp:t>", xml)
                    parts.extend(m for m in matches if m.strip())
                except Exception:
                    continue
    return "\n".join(parts)


def _from_xlsx(path: str) -> str:
    """엑셀 파일의 모든 시트를 읽어 '시트명 | 헤더 | row' 형태로 정리."""
    from openpyxl import load_workbook  # type: ignore
    wb = load_workbook(filename=path, data_only=True, read_only=True)
    parts = []
    try:
        for sheet in wb.worksheets:
            parts.append(f"## 시트: {sheet.title}")
            rows_out = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                # 완전 빈 행 건너뛰기
                if not any(cells):
                    continue
                rows_out.append(" | ".join(cells))
                # 시트당 최대 500행까지만 (너무 큰 엑셀 폭주 방지)
                if len(rows_out) >= 500:
                    rows_out.append("... (이하 생략)")
                    break
            parts.extend(rows_out)
            parts.append("")
    finally:
        wb.close()
    return "\n".join(parts)


def _from_csv(path: str) -> str:
    """CSV를 표 형태 텍스트로 변환."""
    import csv
    raw = _from_text(path)  # 인코딩 감지
    reader = csv.reader(raw.splitlines())
    parts = []
    for i, row in enumerate(reader):
        cells = [c.strip() for c in row]
        if any(cells):
            parts.append(" | ".join(cells))
        if i >= 1000:
            parts.append("... (이하 생략)")
            break
    return "\n".join(parts)


def _from_text(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    # 인코딩 감지
    encoding = "utf-8"
    try:
        from charset_normalizer import from_bytes  # type: ignore
        result = from_bytes(raw).best()
        if result:
            encoding = result.encoding or "utf-8"
    except Exception:
        pass
    try:
        return raw.decode(encoding, errors="ignore")
    except Exception:
        return raw.decode("utf-8", errors="ignore")


def _clean_whitespace(text: str) -> str:
    if not text:
        return ""
    # 연속 공백/줄바꿈 정리
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
